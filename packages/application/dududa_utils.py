# -*- coding: utf-8 -*-
"""Phase 4 拆分：插件纯函数与常量（平台无关，无 Main 依赖）。

原 main.py 模块级函数与常量原样迁移；main.py 保留 re-export 以兼容测试。
"""
import os
import re
import json as _json
import logging
from io import BytesIO

from packages.safeguards.security import Redactor

logger = logging.getLogger("dududa20")

_IMAGE_EXTS = {"jpg", "jpeg", "png", "gif", "webp", "bmp"}

_IGNORE_PATTERNS = {
    "ok", "OK", "Ok", "嗯", "好", "好的", "行", "可以", "对",
    "1", "2", "3", "0", "是", "是的", "对呀", "没错", "哦", "噢",
    "哈哈", "嘿嘿", "呵呵", "。。。", "……", ".....", "？", "?", "！",
}

# Restricted 数据（文档 2.5.9）：密码/Token/Cookie/私钥/QQ 登录态
# 不进 Memory、不发模型或 Tool。
_RESTRICTED_PATTERNS = (
    re.compile(r"sk-[A-Za-z0-9_-]{16,}"),
    re.compile(r"Bearer\s+[A-Za-z0-9._~+/=-]{16,}", re.IGNORECASE),
    re.compile(r"-----BEGIN [A-Z ]*PRIVATE KEY-----"),
    re.compile(r"(?:password|passwd|pwd|secret|api[_-]?key|access[_-]?token|refresh[_-]?token|session[_-]?id|client[_-]?secret)\s*[:=]\s*\S+", re.IGNORECASE),
    re.compile(r"(?:p_skey|skey|p_uin|pt4_token|clientkey|bkn|qqmusic_key|uin)\s*[:=]\s*[A-Za-z0-9_\-]+", re.IGNORECASE),
    re.compile(r"(?:uin|skey|p_skey|pt4_token|clientkey)\s*=\s*[A-Za-z0-9_\-]+", re.IGNORECASE),
    re.compile(r"cookie\s*[:=]\s*\S+", re.IGNORECASE),
    re.compile(r"(?:密码|口令|密钥|令牌|私钥|登录态)\s*[:：=]\s*\S+"),
)
# Sensitive 数据群聊默认不得返回（课表/成绩/健康/位置/私聊等）
_SENSITIVE_GROUP_KW = ("成绩", "课表", "个人安排", "健康", "位置", "就诊", "体检")

_REDACTOR = Redactor()


def _redact_text(text: str) -> str:
    """统一脱敏：credential / URL user-info / URL query / 嵌套值。"""
    try:
        out, _ = _REDACTOR.redact(text or "")
        return out or ""
    except Exception:
        return text or ""


def _contains_restricted(text: str) -> bool:
    """检测 Restricted 数据（密码/Token/Cookie/私钥/QQ 登录态）。"""
    try:
        return any(p.search(text or "") for p in _RESTRICTED_PATTERNS)
    except Exception:
        return False


def _atomic_write_json(path: str, obj) -> None:
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        _json.dump(obj, f, ensure_ascii=False)
    os.replace(tmp, path)


def _group_safe_observations(observations, is_group: bool) -> tuple:
    """群聊默认不返回 Sensitive 数据（课表/成绩/健康/位置/私聊）。"""
    out = []
    for o in observations or ():
        if not o.success or o.data is None:
            continue
        text = str(o.data)
        if not text.strip() or text.strip() in ("[]", "{}", ""):
            continue
        if is_group and any(kw in text for kw in _SENSITIVE_GROUP_KW):
            continue
        out.append(o)
    return tuple(out)


def _detect_media(event) -> tuple:
    for c in event.get_messages():
        t = str(getattr(c, "type", ""))
        if "File" in t or "Image" in t:
            name = getattr(c, "name", "") or getattr(c, "file_name", "") or "media"
            is_image = "Image" in t
            # QQ official bot: url/file/path can be local or http
            for url_attr in ("url", "file_", "file", "path"):
                url = getattr(c, url_attr, "")
                if url and (url.startswith("http") or url.startswith("/")):
                    return url, name, is_image
    return "", "", False


def _has_media_in_raw(event) -> bool:
    try:
        raw = getattr(event, "raw_message", None)
        if raw is None: return False
        for attr in ("message", "json"):
            msg = getattr(raw, attr, None)
            if callable(msg): msg = msg()
            if isinstance(msg, list):
                for item in msg:
                    if isinstance(item, dict) and item.get("type") in ("file", "image"):
                        return True
    except Exception: pass
    return False


def _file_ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _parse_document(data: bytes, filename: str) -> str | None:
    ext = _file_ext(filename)
    try:
        if ext in ("txt","md","py","json","log","csv","yml","yaml"):
            return data.decode("utf-8", errors="replace")
        if ext == "docx":
            from docx import Document
            doc = Document(BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for r in t.rows:
                    parts.append(" | ".join(c.text for c in r.cells))
            return "\n".join(parts)
        if ext == "pdf":
            from pypdf import PdfReader
            return "\n".join(p.extract_text() or "" for p in PdfReader(BytesIO(data)).pages)
        if ext in ("xlsx","xls"):
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(data), read_only=True)
            parts = []
            for sn in wb.sheetnames:
                ws = wb[sn]; parts.append(f"--- {sn} ---")
                for row in ws.iter_rows(values_only=True):
                    parts.append(" | ".join(str(c or "") for c in row))
            wb.close(); return "\n".join(parts)
        return data.decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning("Parse error [%s]: %s", filename, e)
        return None
