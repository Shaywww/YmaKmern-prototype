# -*- coding: utf-8 -*-
"""Phase 4 拆分：插件纯函数与常量（平台无关，无 Main 依赖）。

原 main.py 模块级函数与常量原样迁移；main.py 保留 re-export 以兼容测试。
"""
import ntpath
import os
import posixpath
import re
import json as _json
import logging
from io import BytesIO
from collections.abc import Mapping

from dududa.safeguards.security import Redactor

from dududa.application.dududa_log import get_logger as _get_logger
logger = _get_logger("dududa20")

_IMAGE_EXTS = {"jpg", "jpeg", "png", "gif", "webp", "bmp"}
_VIDEO_EXTS = {"mp4", "mov", "m4v", "webm", "mkv", "avi", "flv"}

_IGNORE_PATTERNS = {
    "ok", "OK", "Ok", "嗯", "好", "好的", "行", "可以", "对",
    "1", "2", "3", "0", "是", "是的", "对呀", "没错", "哦", "噢",
    "哈哈", "嘿嘿", "呵呵", "。。。", "……", ".....", "？", "?", "！",
}

_TEXT_GREETING_RE = re.compile(
    r"^(?:(?:你好|您好|大家好|各位好|嗨|哈喽|在吗|在嘛|"
    r"早上好|中午好|下午好|晚上好|早安|午安|晚安|拜拜|再见)"
    r"(?:呀|啊|哇|啦|呢|哦|喔)?|(?:hi|hello|hey)(?:\s+there)?|"
    r"哈+|嘿+|呵+|h+|233+)[~～!！。,.，…？?]*$",
    re.IGNORECASE,
)
_EMOJI_GREETING_RE = re.compile(
    r"^[\U0001F300-\U0001FAFF\u2600-\u27BF\uFE0F]+$")


def _is_textual_greeting(text: str) -> bool:
    """Return true only when the whole message is a textual greeting.

    This deliberately uses ``fullmatch`` semantics.  A sentence such as
    ``你为什么不说晚上好`` mentions a greeting but is not itself a greeting.
    """
    t = " ".join(str(text or "").strip().lower().split())
    return bool(t and _TEXT_GREETING_RE.fullmatch(t))

def _is_greeting_text(text: str) -> bool:
    """是否纯问候/轻互动（文档 2.5.4：REACT / greeting_only 判定）。

    短名词（USTC/AI/课程名）不属于问候，避免被归为 REACT 或套话回复。
    """
    t = (text or "").strip().lower()
    if not t:
        return False
    return (_is_textual_greeting(t)
            or _EMOJI_GREETING_RE.fullmatch(t) is not None)

# Restricted 数据（文档 2.5.9）：密码/Token/Cookie/私钥/QQ 登录态
# 不进 Memory、不发模型或 Tool。
_RESTRICTED_PATTERNS = (
    re.compile(r"sk-[A-Za-z0-9_-]{16,}"),
    re.compile(r"Bearer\s+[A-Za-z0-9._~+/=-]{16,}", re.IGNORECASE),
    re.compile(r"-----BEGIN [A-Z ]*PRIVATE KEY-----"),
    re.compile(r"(?:password|passwd|pwd|secret|api[_-]?key|access[_-]?token|refresh[_-]?token|session[_-]?id|client[_-]?secret)\s*[:=]\s*\S+", re.IGNORECASE),
    re.compile(r"(?:p_skey|skey|p_uin|pt4_token|clientkey|bkn|qqmusic_key|uin)\s*[:=]\s*[A-Za-z0-9_\-]+", re.IGNORECASE),
    re.compile(r"(?:uin|skey|p_skey|pt4_token|clientkey)\s*=\s*[A-Za-z0-9_\-]+", re.IGNORECASE),
    re.compile(r"cookie\s*[:=]\s*\S+", re.IGNORECASE),
    re.compile(r"(?:密码|口令|密钥|令牌|私钥|登录态)\s*[:：=]\s*\S+"),
)
# Sensitive 数据群聊默认不得返回（课表/成绩/健康/位置/私聊等）
_SENSITIVE_GROUP_KW = ("成绩", "课表", "个人安排", "健康", "位置", "就诊", "体检")

_REDACTOR = Redactor()


def _redact_text(text: str) -> str:
    """统一脱敏：credential / URL user-info / URL query / 嵌套值。"""
    try:
        out, _ = _REDACTOR.redact(text or "")
        return out or ""
    except Exception:
        return text or ""


def _contains_restricted(text: str) -> bool:
    """检测 Restricted 数据（密码/Token/Cookie/私钥/QQ 登录态）。"""
    try:
        return any(p.search(text or "") for p in _RESTRICTED_PATTERNS)
    except Exception:
        return False


def _atomic_write_json(path: str, obj) -> None:
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        _json.dump(obj, f, ensure_ascii=False)
    os.replace(tmp, path)


def _group_safe_observations(observations, is_group: bool) -> tuple:
    """群聊默认不返回 Sensitive 数据（课表/成绩/健康/位置/私聊）。"""
    out = []
    for o in observations or ():
        if not o.success or o.data is None:
            continue
        text = str(o.data)
        if not text.strip() or text.strip() in ("[]", "{}", ""):
            continue
        if is_group and any(kw in text for kw in _SENSITIVE_GROUP_KW):
            continue
        out.append(o)
    return tuple(out)


def _raw_message_segments(event) -> tuple:
    """Return normalized OneBot message segments without depending on an adapter type.

    AstrBot's aiocqhttp event is mapping-like, while tests and some other adapters
    expose ``message``/``json`` as attributes.  NapCat-specific image metadata is
    intentionally read at this trust boundary instead of patching AstrBot itself.
    """
    raw_candidates = [getattr(event, "raw_message", None)]
    raw_candidates.append(
        getattr(getattr(event, "message_obj", None), "raw_message", None))
    seen = set()
    for raw in raw_candidates:
        if raw is None or id(raw) in seen:
            continue
        seen.add(id(raw))
        values = []
        if isinstance(raw, Mapping):
            values.extend((raw.get("message"), raw.get("json")))
        for attr in ("message", "json"):
            try:
                value = getattr(raw, attr, None)
                if callable(value):
                    value = value()
                values.append(value)
            except Exception:
                continue
        for value in values:
            if isinstance(value, Mapping):
                value = value.get("message")
            if isinstance(value, list):
                return tuple(item for item in value if isinstance(item, Mapping))
    return ()


def _segment_data(segment) -> dict:
    """Flatten a OneBot segment while keeping top-level compatibility fields."""
    if not isinstance(segment, Mapping):
        return {}
    out = dict(segment)
    nested = segment.get("data")
    if isinstance(nested, Mapping):
        out.update(nested)
    return out


_STICKER_SUMMARY_RE = re.compile(
    r"(?:表情包?|动画表情|商城表情|贴纸|sticker|mface|emoji)", re.IGNORECASE)


def _detect_media_kind(event) -> str:
    """Classify explicit platform media signals as ``sticker`` or ``image``.

    NapCat represents QQ market faces as OneBot ``image`` segments, but preserves
    ``emoji_id``/``emoji_package_id``/``key`` and ``summary``.  Numeric
    ``sub_type`` is deliberately not used alone because its meaning differs among
    OneBot implementations.  Ambiguous ordinary image segments are left for the
    vision model to distinguish from user-made meme images.
    """
    saw_image = False
    for segment in _raw_message_segments(event):
        kind = str(segment.get("type", "") or "").lower()
        data = _segment_data(segment)
        if kind in ("video", "shortvideo"):
            return "video"
        if kind in ("face", "mface"):
            return "sticker"
        if kind != "image":
            continue
        saw_image = True
        if any(data.get(key) not in (None, "")
               for key in ("emoji_id", "emoji_package_id", "key")):
            return "sticker"
        summary = str(data.get("summary", "") or "")
        if summary and _STICKER_SUMMARY_RE.search(summary):
            return "sticker"
        subtype = str(data.get("sub_type", "") or "").lower()
        if subtype in ("face", "mface", "sticker", "emoji"):
            return "sticker"
        media_name = str(data.get("name", "") or data.get("file_name", "")
                         or data.get("file", "") or data.get("url", ""))
        if _file_ext(media_name.split("?", 1)[0]) == "gif":
            return "gif"
    try:
        for component in event.get_messages() or ():
            type_name = str(getattr(component, "type", "") or "").lower()
            if "video" in type_name:
                return "video"
    except Exception:
        pass
    return "image" if saw_image else "unknown"


def _detect_media(event) -> tuple:
    try:
        components = event.get_messages() or ()
    except Exception:
        components = ()
    for c in components:
        t = str(getattr(c, "type", ""))
        if "File" in t or "Image" in t or "Video" in t:
            name = getattr(c, "name", "") or getattr(c, "file_name", "") or "media"
            is_image = "Image" in t
            # QQ official bot: url/file/path can be local or http
            for url_attr in ("url", "file_", "file", "path"):
                url = getattr(c, url_attr, "")
                if _is_media_source(url):
                    return url, name, is_image
    # Some AstrBot adapters discard mface or cannot construct an Image component
    # containing NapCat extension fields.  The raw OneBot segment still has a URL.
    for segment in _raw_message_segments(event):
        kind = str(segment.get("type", "") or "").lower()
        if kind not in ("image", "mface", "file", "video", "shortvideo"):
            continue
        data = _segment_data(segment)
        is_image = kind in ("image", "mface")
        url = str(data.get("url", "") or "")
        file_value = str(data.get("file", "") or "")
        if not _is_media_source(url):
            url = file_value
        if not _is_media_source(url):
            continue
        name = str(data.get("name", "") or data.get("file_name", "") or "")
        if not name:
            # NapCat's file field is frequently an opaque id, but can also be a path.
            name = os.path.basename(file_value.split("?", 1)[0]) or (
                "image" if is_image else (
                    "video.mp4" if kind in ("video", "shortvideo")
                    else "media"))
        return url, name, is_image
    return "", "", False


def _is_local_media_path(value) -> bool:
    """Recognise POSIX and Windows absolute paths on every host OS.

    ``os.path.isabs`` follows the host platform, so a Linux CI runner does not
    recognise ``C:\\...`` and Windows does not recognise every POSIX fixture.
    Media metadata can originate on a different host from the runtime; accept
    both lexical forms and let the later existence check decide availability.
    """
    if not isinstance(value, str) or not value.strip():
        return False
    path = value.strip()
    return (os.path.isabs(path) or ntpath.isabs(path)
            or posixpath.isabs(path))


def _is_media_source(value) -> bool:
    """Return whether *value* is a supported remote, inline or local source."""
    if not isinstance(value, str):
        return False
    source = value.strip()
    return bool(
        source.startswith(("http://", "https://", "data:"))
        or _is_local_media_path(source)
    )


def _has_media_in_raw(event) -> bool:
    try:
        return any(str(item.get("type", "")).lower()
                   in ("file", "image", "mface", "face", "video", "shortvideo")
                   for item in _raw_message_segments(event))
    except Exception: pass
    return False


def _file_ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _parse_document(data: bytes, filename: str) -> str | None:
    ext = _file_ext(filename)
    try:
        if ext in ("txt","md","py","json","log","csv","yml","yaml"):
            return data.decode("utf-8", errors="replace")
        if ext == "docx":
            from docx import Document
            doc = Document(BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for r in t.rows:
                    parts.append(" | ".join(c.text for c in r.cells))
            return "\n".join(parts)
        if ext == "pdf":
            from pypdf import PdfReader
            return "\n".join(p.extract_text() or "" for p in PdfReader(BytesIO(data)).pages)
        if ext in ("xlsx","xls"):
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(data), read_only=True)
            parts = []
            for sn in wb.sheetnames:
                ws = wb[sn]; parts.append(f"--- {sn} ---")
                for row in ws.iter_rows(values_only=True):
                    parts.append(" | ".join(str(c or "") for c in row))
            wb.close(); return "\n".join(parts)
        return data.decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning("Parse error [%s]: %s", filename, e)
        return None
